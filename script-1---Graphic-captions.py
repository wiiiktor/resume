from pprint import pprint
from pydantic import BaseModel
import os
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains.llm import LLMChain
import docx
from dotenv import load_dotenv, find_dotenv
load_dotenv(find_dotenv())

def read_docx(filename: str):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

def get_files_from_dir(dir_path: str):
    res = []
    # Iterate directory
    for file_path in os.listdir(dir_path):
        # check if current file_path is a file
        if os.path.isfile(os.path.join(dir_path, file_path)):
            # add filename to list
            res.append(dir_path + file_path)
    return res

class RawDocument(BaseModel):
    page_content: str
    metadata: dict

def convert_text_to_page_content_format(text):
    return RawDocument(page_content=text, metadata={})


DOCX_PATH = './files/docx/'

list_of_files = get_files_from_dir(DOCX_PATH)
pprint(list_of_files, width = 200)

chat = ChatOpenAI(temperature=0, model_name="gpt-4", max_tokens=700)

text = read_docx(list_of_files[0])
print(f"Długość załadowanego pliku w znakach: {len(text)}, w słowach: {len(text.split())}")
text = convert_text_to_page_content_format(text)


if isinstance(text, RawDocument):
    text = text.page_content

num_graphics = text.count("graphic-number")

template = """
In a document you will find {num_graphics} codes in a format 
graphic-number-xxx where xxx are three integers.
For example graphic-number-003.
Your aim is to make a brief summary of the text around the codes, 
especially in a paragraph just before the text.
You provide a reply in a format:
    ("graphic-number-001": "description to the graphic")

Document: {document}
"""

prompt = PromptTemplate(
    input_variables = ["num_graphics", "document"],
    template = template
)

chain = LLMChain(llm = chat, prompt = prompt)
captions = chain.run(document = text, num_graphics = num_graphics)
pprint(captions, width=150)